import streamlit as st
from openai import OpenAI
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# DEBUG_MODE = True si tu veux modifier l'appli sans utiliser ChatGPT à chaque fois (Tokens)
DEBUG_MODE = False
FONT_PATH = "DejaVuSans.ttf"

# Setup OpenAI
client = OpenAI(api_key=st.secrets["OPENAI_KEY"])

# Configuration App Streamlit
st.title("🤝 Aide à la candidature")

if st.button("🔄 Recommencer"):
    st.session_state.clear()
    st.rerun()

# Contrôle la progression de l'appli et garde les inputs de l'utilisateur
if "step" not in st.session_state:
    st.session_state.step = "input_mode"
if "user_data" not in st.session_state:
    st.session_state.user_data = {}
if "recommendations" not in st.session_state:
    st.session_state.recommendations = []
if "accepted_suggestions" not in st.session_state:
    st.session_state.accepted_suggestions = []

# Step 1: Choice of input mode
# Choix pour le user: insérer un texte qui le décrit en total, ou répondre à chaque bloc pour être guidé
if st.session_state.step == "input_mode":
    mode = st.radio("Souhaitez-vous entrer un résumé ou remplir les informations une par une ?", ["Résumé global", "Questions une par une"])
    st.session_state.input_mode = mode
    if st.button("Continuer"):
        st.session_state.step = "summary_input" if mode == "Résumé global" else "form_input"

# Step 2A: Il soumets un Résumé global
if st.session_state.step == "summary_input":
    default_summary = ("Je suis motivée, ponctuelle et organisée. J’ai obtenu un CAP Cuisine "
                       "et j’ai travaillé deux ans comme serveuse dans un restaurant local. "
                       "Je suis à l’aise avec le contact client et je sais gérer la caisse. "
                       "J’aimerais trouver un emploi stable dans la restauration ou l’accueil.")
    summary = st.text_area("Écrivez votre résumé ici (en français)", height=250, value=default_summary if DEBUG_MODE else "")
    if st.button("Analyser mon profil"):
        st.session_state.user_data = {"summary": summary.strip()}
        st.session_state.step = "recommend"

# Step 2B: Il utilise les Champs classiques
if st.session_state.step == "form_input":
    with st.form("profile_form"):
        first_name = st.text_input("Prénom", "Jeanne" if DEBUG_MODE else "")
        last_name = st.text_input("Nom", "Dupont" if DEBUG_MODE else "")
        age = st.number_input("Âge", min_value=0, max_value=120, value=32 if DEBUG_MODE else 0)
        location = st.text_input("Ville", "Marseille" if DEBUG_MODE else "")
        description = st.text_area("Décrivez-vous", "Dynamique et organisée, je cherche un emploi stable." if DEBUG_MODE else "")
        education = st.text_area("Votre parcours scolaire", "CAP Cuisine" if DEBUG_MODE else "")
        skills = st.text_area("Vos compétences", "Ponctualité, gestion de caisse, service client" if DEBUG_MODE else "")
        experience = st.text_area("Vos expériences professionnelles", "Serveuse (2 ans), cantine scolaire (1 an)" if DEBUG_MODE else "")
        submitted = st.form_submit_button("Analyser mon profil")

    if submitted:
        st.session_state.user_data = {
            "first_name": first_name.strip(),
            "last_name": last_name.strip(),
            "age": age,
            "location": location.strip(),
            "description": description.strip(),
            "education": education.strip(),
            "skills": skills.strip(),
            "experience": experience.strip()
        }
        st.session_state.step = "recommend"

# Step 3: Basé sur l'input de l'utilisateur, GPT génère 5 simples recommandations de contenu qui match avec le profil
# Les recommendations acceptées sont ajoutées au profil, les autres sont oubliées
if st.session_state.step == "recommend":
    user = st.session_state.user_data

    content = user.get("summary", "") or f"""
Nom: {user.get('first_name', '')} {user.get('last_name', '')}
Âge: {user.get('age', '')}
Lieu: {user.get('location', '')}
Description: {user.get('description', '')}
Éducation: {user.get('education', '')}
Compétences: {user.get('skills', '')}
Expérience: {user.get('experience', '')}
"""

    if not st.session_state.recommendations and not DEBUG_MODE:
        with st.spinner("Analyse de votre profil pour suggestions..."):
            prompt = f"""
Tu es un assistant bienveillant. Voici un profil :
{content}

Propose 5 ajouts logiques et pertinents qui pourraient améliorer ce profil.
Exprime chaque proposition comme une question très simple. Exemple :
- Puis-je ajouter que vous avez travaillé avec des enfants ?
- Puis-je ajouter que vous parlez anglais ?
"""
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}]
            )
            st.session_state.recommendations = [
                s.strip("-• ").strip() for s in response.choices[0].message.content.strip().split("\n") if s.strip()
            ]
    elif DEBUG_MODE:
        st.session_state.recommendations = [
            "Puis-je ajouter que vous parlez anglais ?",
            "Puis-je ajouter que vous avez le permis B ?",
            "Puis-je ajouter que vous savez utiliser Word et Excel ?",
            "Puis-je ajouter que vous êtes à l’aise en équipe ?",
            "Puis-je ajouter que vous avez travaillé avec des enfants ?"
        ]

    for i, rec in enumerate(st.session_state.recommendations):
        if rec:
            col1, col2, col3 = st.columns([5, 1, 1])
            with col1: st.write(f"👉 {rec}")
            with col2:
                if st.button("✅", key=f"accept_{i}"):
                    st.session_state.accepted_suggestions.append(rec)
                    st.session_state.recommendations[i] = None
                    st.rerun()
            with col3:
                if st.button("❌", key=f"reject_{i}"):
                    st.session_state.recommendations[i] = None
                    st.rerun()

    if all(r is None for r in st.session_state.recommendations):
        if st.button("➡️ Continuer vers la génération du PDF"):
            st.session_state.step = "generate"

# Step 4: DOCX generation
if st.session_state.step == "generate":
    user = st.session_state.user_data
    st.subheader("📝 Résultat final")

    sections = []
    if user.get("description"): sections.append(f"🧍 Description :\n{user['description']}")
    if user.get("education"): sections.append(f"🎓 Éducation :\n{user['education']}")
    if user.get("skills"): sections.append(f"🛠️ Compétences :\n{user['skills']}")
    if user.get("experience"): sections.append(f"💼 Expérience :\n{user['experience']}")
    if st.session_state.accepted_suggestions:
        sections.append("➕ Informations ajoutées :\n- " + "\n- ".join(st.session_state.accepted_suggestions))

    profile_input = f"""
    Description: {user.get('description', '')}
    Éducation: {user.get('education', '')}
    Compétences: {user.get('skills', '')}
    Expérience: {user.get('experience', '')}
    Suggestions supplémentaires: {' | '.join(st.session_state.accepted_suggestions or [])}
    """

    reformulate_prompt = f"""
Tu es un assistant RH bienveillant. À partir des informations suivantes, rédige un contenu clair, professionnel et valorisant, structuré comme un CV.

Organise la sortie finale en 4 sections bien rédigées :
- Une **description** du profil (2-4 phrases maximum)
- Une section **Éducation** (études, diplômes)
- Une section **Compétences** (sous forme de liste claire)
- Une section **Expérience** (avec missions ou tâches, 2-3 lignes max par poste)

Voici les informations à traiter :

{profile_input}

✅ Rédige directement le texte final, sans poser de questions.
✅ Si certaines suggestions sont utiles, intègre-les naturellement au bon endroit.
✅ Ne reformule pas les titres. Commence directement par :
Description :
Éducation :
...
"""


    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Tu aides des personnes à rédiger des CV professionnels à partir de données brutes."},
            {"role": "user", "content": reformulate_prompt}
        ]
    )

    profile_text = response.choices[0].message.content.strip()
    st.text(profile_text)

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    doc.add_heading(f"{user.get('first_name', '')} {user.get('last_name', '')}", level=1)
    def add_section(title, content):
        if content:
            p_title = doc.add_paragraph()
            run_title = p_title.add_run(f"{title} :")
            run_title.bold = True
            doc.add_paragraph("")  # Spacer
            doc.add_paragraph(content)
            doc.add_paragraph("")  # Spacer

    # Break down profile_text into parts (simple logic for now)
    sections = {"Description": "", "Éducation": "", "Compétences": "", "Expérience": ""}
    current_section = None
    for line in profile_text.split("\n"):
        if any(line.strip().startswith(title) for title in sections.keys()):
            current_section = next(title for title in sections.keys() if line.strip().startswith(title))
            sections[current_section] = line.replace(f"{current_section} :", "").strip()
        elif current_section:
            sections[current_section] += " " + line.strip()

    # Add formatted content
    for title, content in sections.items():
        add_section(title, content)
    doc.save("mon_cv.docx")

    with open("mon_cv.docx", "rb") as f:
        st.download_button("📥 Télécharger le CV (.docx)", f, file_name="mon_cv.docx")

 